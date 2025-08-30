import os
import re
import json
import sqlite3
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Tuple

from fastapi import FastAPI, Request, Form, UploadFile, File, Query, Response
from fastapi.responses import RedirectResponse, HTMLResponse, FileResponse, PlainTextResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from starlette.middleware.sessions import SessionMiddleware
from starlette import status

# Documenti
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader

# Password hashing (puro Python)
from passlib.hash import pbkdf2_sha256

# OpenAI (opzionale)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
if OPENAI_API_KEY:
    try:
        from openai import OpenAI
        ai_client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception:
        ai_client = None
else:
    ai_client = None

APP_NAME = "VoxUp"
SECRET_KEY = os.getenv("SECRET_KEY", "change-me-please-very-long-secret")

# ----------------------------
# Cartelle dati & DB
# ----------------------------
def get_writable_data_dir() -> str:
    cand = []
    if os.getenv("DATA_DIR"):
        cand.append(os.getenv("DATA_DIR"))
    cand.append("/var/data")  # Render Disk tipico
    cand.append(os.path.join(os.getcwd(), "data"))
    for p in cand:
        try:
            os.makedirs(p, exist_ok=True)
            test = os.path.join(p, ".write_test")
            with open(test, "w", encoding="utf-8") as f:
                f.write("ok")
            os.remove(test)
            return p
        except Exception:
            continue
    return os.getcwd()

DATA_DIR = get_writable_data_dir()
DB_PATH = os.path.join(DATA_DIR, "voxup.sqlite3")
DRAFTS_PATH = os.path.join(DATA_DIR, "voxup_drafts.json")
NOTES_PATH = os.path.join(DATA_DIR, "voxup_notes.json")
NEWS_CACHE_PATH = os.path.join(DATA_DIR, "voxup_news_cache.json")

for _p, _default in [(DRAFTS_PATH, []), (NOTES_PATH, [])]:
    if not os.path.exists(_p):
        try:
            with open(_p, "w", encoding="utf-8") as f:
                json.dump(_default, f)
        except Exception:
            pass

def db_connect() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def db_init():
    conn = db_connect()
    cur = conn.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        email TEXT NOT NULL UNIQUE,
        role TEXT DEFAULT '',
        password_hash TEXT NOT NULL,
        created_at TEXT NOT NULL
    );
    """)
    conn.commit()
    conn.close()
db_init()

# ----------------------------
# App & Templates
# ----------------------------
app = FastAPI(title=APP_NAME)
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# ----------------------------
# Helpers - formattazione & util
# ----------------------------
def ensure_session_defaults(session: Dict[str, Any]) -> None:
    session.setdefault("auth", False)
    session.setdefault("user", None)  # può essere None
    session.setdefault("onboarding_done", False)
    session.setdefault("profile", {
        "first_name": "",
        "last_name": "",
        "role": "",
        "tones": [],
        "tone_other": "",
        "channels": ["Social"],
        "add_ai": False
    })
    session.setdefault("style_guide", "")
    session.setdefault("last_results", {})

def unicode_bold(text: str) -> str:
    def _bold_char(c: str) -> str:
        if 'A' <= c <= 'Z': return chr(ord('𝐀') + (ord(c) - ord('A')))
        if 'a' <= c <= 'z': return chr(ord('𝐚') + (ord(c) - ord('a')))
        if '0' <= c <= '9': return chr(ord('𝟎') + (ord(c) - ord('0')))
        return c
    return "".join(_bold_char(c) for c in text)

EMOJI_PATTERN = re.compile("["                     
    "\U0001F600-\U0001F64F"
    "\U0001F300-\U0001F5FF"
    "\U0001F680-\U0001F6FF"
    "\U0001F1E0-\U0001F1FF"
    "\u2600-\u26FF"
    "\u2700-\u27BF"
    "]+", flags=re.UNICODE)

def remove_emojis(text: str) -> str:
    return EMOJI_PATTERN.sub("", text)

def format_for_channel(base_text: str, channel: str) -> str:
    """Social = bold Unicode; Sito/Stampa = <strong> senza emoji."""
    if channel.lower() == "social":
        return unicode_bold(base_text)
    else:
        no_emoji = remove_emojis(base_text)
        if ":" in no_emoji:
            head, tail = no_emoji.split(":", 1)
            return f"<strong>{head.strip()}:</strong>{tail}"
        return f"<strong>{no_emoji}</strong>"

def extract_text_from_upload(filename: str, data: bytes) -> str:
    """Supporta .txt/.md, .pdf, .docx (estrazione semplice)."""
    name = filename.lower()
    if name.endswith((".txt", ".md")):
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    if name.endswith(".pdf"):
        try:
            reader = PdfReader(BytesIO(data))
            parts = []
            for page in reader.pages[:20]:
                text = page.extract_text() or ""
                parts.append(text)
            return "\n".join(parts).strip()
        except Exception:
            return ""
    if name.endswith(".docx"):
        try:
            doc = Document(BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    return ""

def split_into_posts(text: str, limit: int = 280) -> List[str]:
    text = text.strip()
    if not text: return []
    words = text.split()
    chunks, cur = [], ""
    for w in words:
        candidate = (cur + " " + w).strip() if cur else w
        if len(candidate) <= limit:
            cur = candidate
        else:
            if cur: chunks.append(cur)
            if len(w) > limit:
                while len(w) > limit:
                    chunks.append(w[:limit]); w = w[limit:]
                cur = w
            else:
                cur = w
    if cur: chunks.append(cur)
    if len(chunks) > 1:
        total = len(chunks)
        out = []
        for i, c in enumerate(chunks, 1):
            prefix = f"{i}/{total} "
            room = limit - len(prefix)
            out.append(prefix + (c[:room] if len(c) > room else c))
        return out
    return chunks

def save_draft(entry: Dict[str, Any]) -> None:
    """Aggiunge una bozza su file JSON, max 50 elementi."""
    try:
        with open(DRAFTS_PATH, "r", encoding="utf-8") as f:
            arr = json.load(f)
    except Exception:
        arr = []
    arr.insert(0, entry)
    arr = arr[:50]
    try:
        with open(DRAFTS_PATH, "w", encoding="utf-8") as f:
            json.dump(arr, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

# ----------------------------
# AI: rielaborazione vera (con fallback)
# ----------------------------
def build_ai_prompt(base_context: str, profile: Dict[str, Any], style_guide: str) -> Tuple[str, str]:
    full_name = f"{profile.get('first_name','').strip()} {profile.get('last_name','').strip()}".strip()
    role = profile.get("role","").strip()
    tones = profile.get("tones", [])
    tone_other = profile.get("tone_other","").strip()
    tono_testo = ", ".join([t for t in tones if t] + ([tone_other] if tone_other else [])) or "istituzionale"
    sys = (
        "Sei un addetto stampa politico italiano. "
        "Rielabora i contenuti in un testo chiaro, sintetico e coerente con il ruolo indicato. "
        "Non inventare fatti. Mantieni un tono adatto ai 'toni' richiesti."
    )
    usr = (
        f"Ruolo: {role}\n"
        f"Nome: {full_name}\n"
        f"Toni: {tono_testo}\n"
        f"Stile guida (estratto, se presente): {style_guide[:600]}\n\n"
        f"Contesto da rielaborare:\n{base_context[:6000]}"
        "\n\nScrivi un comunicato breve di 6-10 frasi con una citazione in prima persona del politico."
    )
    return sys, usr

def ai_rewrite_or_fallback(base_context: str, profile: Dict[str, Any], style_guide: str) -> str:
    """Tenta rielaborazione con OpenAI; se non disponibile, fallback deterministico."""
    try:
        if ai_client:
            sys, usr = build_ai_prompt(base_context, profile, style_guide)
            comp = ai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role":"system","content":sys},
                    {"role":"user","content":usr}
                ],
                temperature=0.5,
                max_tokens=700
            )
            out = comp.choices[0].message.content.strip()
            return out
    except Exception:
        pass
    # Fallback semplice e sicuro (nessuna chiamata esterna)
    full_name = f"{profile.get('first_name','').strip()} {profile.get('last_name','').strip()}".strip()
    role = profile.get("role","").strip()
    return (
        f"{role} {full_name}: ecco i punti principali.\n"
        f"- {base_context[:240]}...\n\n"
        "Dichiarazione: \"Mettiamo al centro cittadini e territori. "
        "Lavoriamo con serietà e risultati concreti.\""
    )

def require_auth(request: Request) -> bool:
    ensure_session_defaults(request.session)
    return bool(request.session.get("auth"))

# ----------------------------
# NEWS: fonti RSS e cache
# ----------------------------
FEEDS: Dict[str, str] = {
    # Italia
    "ANSA": "https://www.ansa.it/sito/ansait_rss.xml",
    "Repubblica": "https://www.repubblica.it/rss/homepage/rss2.0.xml",
    "Corriere": "https://xml2.corriereobjects.it/rss/homepage.xml",
    "Il Sole 24 Ore": "https://www.ilsole24ore.com/rss/italia.xml",
    "AGI": "https://www.agi.it/rss/ultime-notizie.xml",
    # Internazionali
    "BBC": "http://feeds.bbci.co.uk/news/rss.xml",
    "Reuters": "http://feeds.reuters.com/reuters/topNews",
    "AP": "https://apnews.com/hub/apf-topnews?utm_source=ap_rss&utm_medium=rss&utm_campaign=ap_rss",
    "The Guardian": "https://www.theguardian.com/world/rss",
    "Politico EU": "https://www.politico.eu/feed/"
}

NEWS_TTL_MIN = 15

def load_news_from_cache() -> Tuple[List[Dict[str, Any]], datetime]:
    try:
        with open(NEWS_CACHE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("items", []), datetime.fromisoformat(data.get("ts"))
    except Exception:
        return [], datetime.min

def save_news_cache(items: List[Dict[str, Any]]) -> None:
    try:
        with open(NEWS_CACHE_PATH, "w", encoding="utf-8") as f:
            json.dump({"ts": datetime.utcnow().isoformat(), "items": items}, f, ensure_ascii=False)
    except Exception:
        pass

def fetch_feeds() -> List[Dict[str, Any]]:
    try:
        import feedparser
    except Exception:
        return []
    items: List[Dict[str, Any]] = []
    for source, url in FEEDS.items():
        try:
            feed = feedparser.parse(url)
            for e in feed.entries[:10]:
                title = getattr(e, "title", "").strip()
                link = getattr(e, "link", "")
                published = getattr(e, "published", "") or getattr(e, "updated", "")
                if title and link:
                    items.append({"title": title, "link": link, "source": source, "published": published})
        except Exception:
            continue
    # Dedup by title
    dedup = {}
    for it in items:
        key = (it["title"][:140] + it["source"])
        if key not in dedup:
            dedup[key] = it
    # keep top ~60
    return list(dedup.values())[:60]

def get_news_items() -> List[Dict[str, Any]]:
    cached, ts = load_news_from_cache()
    if ts != datetime.min and datetime.utcnow() - ts < timedelta(minutes=NEWS_TTL_MIN) and cached:
        return cached
    items = fetch_feeds()
    if items:
        save_news_cache(items)
        return items
    return cached

# ----------------------------
# ROUTES: health + HEAD
# ----------------------------
@app.head("/")
def head_root():
    return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)

@app.head("/health")
def head_health():
    return Response(status_code=200)

@app.get("/health")
def health():
    return {"status": "ok", "app": APP_NAME, "data_dir": DATA_DIR}

# ----------------------------
# AUTH: register / login / logout
# ----------------------------
@app.get("/register", response_class=HTMLResponse)
def register_page(request: Request):
    ensure_session_defaults(request.session)
    if request.session.get("auth"):
        return RedirectResponse(url="/", status_code=status.HTTP_302_FOUND)
    return templates.TemplateResponse("register.html", {"request": request, "app_name": APP_NAME})

@app.post("/register")
def register_submit(
    request: Request,
    name: str = Form(...),
    email: str = Form(...),
    role: str = Form(""),
    password: str = Form(...),
    password2: str = Form(...)
):
    ensure_session_defaults(request.session)
    if password != password2 or len(password) < 6:
        return templates.TemplateResponse(
            "register.html",
            {"request": request, "app_name": APP_NAME,
             "error": "Password non valida (min 6 caratteri) o non coincidono."},
            status_code=400
        )
    email = email.strip().lower()
    name = name.strip()
    role = role.strip()
    pw_hash = pbkdf2_sha256.hash(password)

    try:
        conn = db_connect()
        cur = conn.cursor()
        cur.execute("INSERT INTO users(name,email,role,password_hash,created_at) VALUES(?,?,?,?,?)",
                    (name, email, role, pw_hash, datetime.utcnow().isoformat()+"Z"))
        conn.commit()
        conn.close()
    except sqlite3.IntegrityError:
        return templates.TemplateResponse(
            "register.html",
            {"request": request, "app_name": APP_NAME,
             "error": "Email già registrata."},
            status_code=400
        )

    # Auto-login dopo registrazione
    request.session["auth"] = True
    request.session["user"] = {"name": name, "email": email, "role": role}
    request.session["profile"] = {
        "first_name": name.split(" ")[0] if name else "",
        "last_name": " ".join(name.split(" ")[1:]) if len(name.split(" ")) > 1 else "",
        "role": role,
        "tones": [],
        "tone_other": "",
        "channels": ["Social"],
        "add_ai": False
    }
    return RedirectResponse(url="/", status_code=status.HTTP_302_FOUND)

@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    ensure_session_defaults(request.session)
    if request.session.get("auth"):
        return RedirectResponse(url="/", status_code=status.HTTP_302_FOUND)
    return templates.TemplateResponse("login.html", {"request": request, "app_name": APP_NAME})

@app.post("/login")
def login_submit(request: Request, email: str = Form(...), password: str = Form(...)):
    ensure_session_defaults(request.session)
    email = email.strip().lower()
    conn = db_connect()
    cur = conn.cursor()
    cur.execute("SELECT id,name,email,role,password_hash FROM users WHERE email=?", (email,))
    row = cur.fetchone()
    conn.close()
    if not row or not pbkdf2_sha256.verify(password, row["password_hash"]):
        return templates.TemplateResponse(
            "login.html",
            {"request": request, "app_name": APP_NAME, "error": "Credenziali non valide."},
            status_code=401
        )

    request.session["auth"] = True
    request.session["user"] = {"name": row["name"], "email": row["email"], "role": row["role"]}
    prof = request.session.get("profile", {})
    if not prof.get("role"):
        prof["role"] = row["role"] or ""
    request.session["profile"] = prof
    return RedirectResponse(url="/", status_code=status.HTTP_302_FOUND)

@app.get("/logout")
def logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)

# ----------------------------
# Onboarding / Pagine app
# ----------------------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    # Se onboarding completato, non mostrare più il form: vai a compose
    if request.session.get("onboarding_done"):
        return RedirectResponse(url="/compose", status_code=status.HTTP_302_FOUND)
    return templates.TemplateResponse("home.html", {
        "request": request,
        "app_name": APP_NAME,
        "profile": request.session.get("profile", {}),
        "user": request.session.get("user") or {}
    })

@app.post("/save_onboarding")
def save_onboarding(
    request: Request,
    first_name: str = Form(""),
    last_name: str = Form(""),
    role: str = Form(""),
    tones: List[str] = Form([]),
    channels: List[str] = Form(["Social"]),
    tone_other: str = Form("")
):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    ensure_session_defaults(request.session)
    profile = request.session["profile"]
    profile["first_name"] = first_name.strip()
    profile["last_name"] = last_name.strip()
    profile["role"] = role.strip()
    profile["tones"] = tones
    profile["channels"] = channels or ["Social"]
    profile["tone_other"] = tone_other.strip()
    request.session["profile"] = profile
    request.session["onboarding_done"] = True
    return RedirectResponse(url="/compose", status_code=status.HTTP_302_FOUND)

@app.get("/compose", response_class=HTMLResponse)
def compose_page(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    return templates.TemplateResponse("compose.html", {
        "request": request,
        "app_name": APP_NAME,
        "profile": request.session.get("profile", {}),
        "results": None,
        "file_previews": [],
        "errors": []
    })

@app.post("/generate", response_class=HTMLResponse)
async def generate(
    request: Request,
    text_input: str = Form(""),
    url_input: str = Form(""),
    split_x: Optional[str] = Form(None),
    files: List[UploadFile] = File(None),
):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)

    ensure_session_defaults(request.session)
    profile = request.session["profile"]
    style_guide = request.session.get("style_guide", "")
    add_ai = profile.get("add_ai", False)

    bodies: List[str] = []
    errors: List[str] = []

    if text_input.strip():
        bodies.append(text_input.strip())
    if url_input.strip():
        bodies.append(f"Fonte: {url_input.strip()}")

    file_previews = []
    if files:
        for f in files:
            try:
                raw = await f.read()
                text = extract_text_from_upload(f.filename, raw)
                snippet = text[:500] + ("…" if len(text) > 500 else "")
                if not text:
                    snippet = "(Anteprima non disponibile: formato non supportato o documento protetto)"
                file_previews.append({"name": f.filename, "snippet": snippet})
                if text:
                    bodies.append(f"File {f.filename}:\n{text[:4000]}")
            except Exception as ex:
                errors.append(f"Errore lettura file {f.filename}: {str(ex)[:120]}")
                file_previews.append({"name": f.filename, "snippet": "(Errore in lettura)"})

    base_context = "\n\n".join(bodies).strip()
    if not base_context:
        base_context = "Nessun testo inserito. Aggiungi un testo, un URL o allega dei file."

    # Punto 2: rielaborazione AI vera (se attiva in Profilo); altrimenti copia “grezza”
    if add_ai:
        rewritten = ai_rewrite_or_fallback(base_context, profile, style_guide)
        working_text = rewritten
    else:
        working_text = base_context

    channels = profile.get("channels", ["Social"])
    results: Dict[str, Any] = {}
    do_split = (split_x == "on")
    for ch in channels:
        if ch == "Social" and do_split:
            pieces = split_into_posts(working_text, 280)
            results[ch] = [unicode_bold(p) for p in pieces]
        else:
            results[ch] = format_for_channel(working_text, ch)

    request.session["last_results"] = results
    save_draft({
        "ts": datetime.utcnow().isoformat() + "Z",
        "profile": request.session.get("profile", {}),
        "input": {"text": text_input, "url": url_input},
        "results": results
    })

    return templates.TemplateResponse("compose.html", {
        "request": request,
        "app_name": APP_NAME,
        "profile": profile,
        "results": results,
        "file_previews": file_previews,
        "split_used": do_split,
        "errors": errors
    })

@app.get("/export")
def export_result(request: Request, channel: str = Query(...), fmt: str = Query(...)):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    ensure_session_defaults(request.session)
    results = request.session.get("last_results", {})
    if not results or channel not in results:
        return PlainTextResponse("Nessun risultato da esportare per questo canale.", status_code=400)

    safe_name = channel.lower()
    value = results[channel]

    if isinstance(value, list):
        joined = "\n\n".join(value)
        if fmt == "txt":
            path = os.path.join(DATA_DIR, f"voxup_{safe_name}.txt")
            with open(path, "w", encoding="utf-8") as f: f.write(joined)
            return FileResponse(path, media_type="text/plain", filename=f"{APP_NAME}_{safe_name}.txt")
        elif fmt == "html":
            html = "<br>".join(value)
            path = os.path.join(DATA_DIR, f"voxup_{safe_name}.html")
            with open(path, "w", encoding="utf-8") as f:
                f.write(f"<!doctype html><meta charset='utf-8'><body>{html}</body>")
            return FileResponse(path, media_type="text/html", filename=f"{APP_NAME}_{safe_name}.html")
        else:
            return PlainTextResponse("Formato non supportato per questo canale.", status_code=400)

    content: str = str(value)
    if fmt == "txt":
        plain = re.sub(r"<[^>]+>", "", content)
        path = os.path.join(DATA_DIR, f"voxup_{safe_name}.txt")
        with open(path, "w", encoding="utf-8") as f: f.write(plain)
        return FileResponse(path, media_type="text/plain", filename=f"{APP_NAME}_{safe_name}.txt")

    if fmt == "html":
        html = content if channel != "Social" else f"<pre>{content}</pre>"
        path = os.path.join(DATA_DIR, f"voxup_{safe_name}.html")
        with open(path, "w", encoding="utf-8") as f:
            f.write(f"<!doctype html><meta charset='utf-8'><body>{html}</body>")
        return FileResponse(path, media_type="text/html", filename=f"{APP_NAME}_{safe_name}.html")

    if fmt == "docx" and channel == "Stampa":
        doc = Document()
        plain = re.sub(r"<[^>]+>", "", content)
        for para in plain.split("\n"): doc.add_paragraph(para)
        tmp = os.path.join(DATA_DIR, f"voxup_{safe_name}.docx")
        doc.save(tmp)
        return FileResponse(tmp,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{APP_NAME}_{safe_name}.docx")

    return PlainTextResponse("Formato non supportato.", status_code=400)

# ----------------------------
# Stile / Profilo
# ----------------------------
@app.get("/style", response_class=HTMLResponse)
def style_page(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    return templates.TemplateResponse("style.html", {
        "request": request, "app_name": APP_NAME,
        "style_guide": request.session.get("style_guide", "")
    })

@app.post("/style")
def style_save(request: Request, style_guide: str = Form("")):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    request.session["style_guide"] = style_guide.strip()
    return RedirectResponse(url="/style", status_code=status.HTTP_302_FOUND)

@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    return templates.TemplateResponse("profile.html", {
        "request": request,
        "app_name": APP_NAME,
        "profile": request.session.get("profile", {}),
        "user": request.session.get("user") or {}
    })

@app.post("/profile")
def profile_save(
    request: Request,
    first_name: str = Form(""),
    last_name: str = Form(""),
    role: str = Form(""),
    tones: List[str] = Form([]),
    channels: List[str] = Form(["Social"]),
    add_ai: Optional[str] = Form(None),
    tone_other: str = Form("")
):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    ensure_session_defaults(request.session)
    profile = request.session["profile"]
    profile["first_name"] = first_name.strip()
    profile["last_name"] = last_name.strip()
    profile["role"] = role.strip()
    profile["tones"] = tones
    profile["channels"] = channels or ["Social"]
    profile["add_ai"] = (add_ai == "on")
    profile["tone_other"] = tone_other.strip()
    request.session["profile"] = profile
    return RedirectResponse(url="/profile", status_code=status.HTTP_302_FOUND)

# ----------------------------
# Bozze
# ----------------------------
@app.get("/drafts", response_class=HTMLResponse)
def drafts_page(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    try:
        with open(DRAFTS_PATH, "r", encoding="utf-8") as f:
            items = json.load(f)
    except Exception:
        items = []
    return templates.TemplateResponse("drafts.html", {
        "request": request, "app_name": APP_NAME, "drafts": items[:20]
    })

# ----------------------------
# News
# ----------------------------
@app.get("/news.json")
def news_json():
    items = get_news_items()
    return JSONResponse(items)

@app.get("/news", response_class=HTMLResponse)
def news_page(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    items = get_news_items()
    return templates.TemplateResponse("news.html", {
        "request": request, "app_name": APP_NAME, "items": items
    })

# ----------------------------
# Blocco Note
# ----------------------------
@app.get("/notes", response_class=HTMLResponse)
def notes_page(request: Request):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    try:
        with open(NOTES_PATH, "r", encoding="utf-8") as f:
            notes = json.load(f)
    except Exception:
        notes = []
    return templates.TemplateResponse("notes.html", {
        "request": request, "app_name": APP_NAME, "notes": notes
    })

@app.post("/notes/add")
def notes_add(request: Request, title: str = Form(""), body: str = Form("")):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    note = {
        "id": int(datetime.utcnow().timestamp()*1000),
        "title": title.strip() or "Senza titolo",
        "body": body.strip(),
        "ts": datetime.utcnow().isoformat()+"Z"
    }
    try:
        with open(NOTES_PATH, "r", encoding="utf-8") as f:
            notes = json.load(f)
    except Exception:
        notes = []
    notes.insert(0, note)
    with open(NOTES_PATH, "w", encoding="utf-8") as f:
        json.dump(notes, f, ensure_ascii=False, indent=2)
    return RedirectResponse(url="/notes", status_code=status.HTTP_302_FOUND)

@app.post("/notes/delete")
def notes_delete(request: Request, note_id: str = Form(...)):
    if not require_auth(request):
        return RedirectResponse(url="/login", status_code=status.HTTP_302_FOUND)
    try:
        with open(NOTES_PATH, "r", encoding="utf-8") as f:
            notes = json.load(f)
    except Exception:
        notes = []
    notes = [n for n in notes if str(n.get("id")) != str(note_id)]
    with open(NOTES_PATH, "w", encoding="utf-8") as f:
        json.dump(notes, f, ensure_ascii=False, indent=2)
    return RedirectResponse(url="/notes", status_code=status.HTTP_302_FOUND)
